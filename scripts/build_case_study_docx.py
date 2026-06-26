from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageOps


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "case-study" / "Foxman_Case_Study.docx"


INK = RGBColor(22, 19, 21)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
GOLD = RGBColor(184, 138, 59)
GREEN = RGBColor(86, 123, 56)
RED = RGBColor(155, 28, 28)
LIGHT = "F4F6F9"
HEADER = "E8EEF5"
PREVIEW_MAX_EDGE = 1800


@dataclass(frozen=True)
class Asset:
    title: str
    path: str
    dimensions: str
    size: str
    status: str
    role: str
    note: str


ASSETS: list[Asset] = [
    Asset("Foxman Atlas", "assets/game/atlases/characters/foxman/foxman_atlas.png", "1274x780", "194,096 bytes", "Packed atlas", "Player frames", "Strong identity; feel still depends on input, animation timing, and camera."),
    Asset("Drunken Guard Atlas", "assets/game/atlases/enemies/drunken_guard/drunken_guard_atlas.png", "1004x967", "235,590 bytes", "Packed atlas", "First enemy", "Readable character idea; encounter pacing remained rough."),
    Asset("Tax Clerk Atlas", "assets/game/atlases/enemies/tax_clerk/tax_clerk_atlas.png", "910x1282", "227,274 bytes", "Packed atlas", "Second-path enemy", "One asset was reused for clerk and elite roles."),
    Asset("Toll Baron Atlas", "assets/game/atlases/enemies/toll_baron/toll_baron_atlas.png", "1003x1269", "405,208 bytes", "Packed atlas", "Mini-boss", "One of the strongest assets; boss art did not automatically produce boss feel."),
    Asset("Pickup And Exit Atlas", "assets/game/atlases/props/rotten_borough/pickup_exit_atlas.png", "882x959", "216,529 bytes", "Packed atlas", "Gate and pickups", "The unlocked gate overpowered the first-room composition until scaled down."),
    Asset("Reward Shop Icons Atlas", "assets/game/atlases/ui/reward_shop_icons/reward_shop_icons_atlas.png", "1278x1278", "446,028 bytes", "Packed atlas", "Shop categories", "Good UI proof, though expensive for the runtime value."),
    Asset("Rotten Borough Runtime Background", "assets/game/backgrounds/rotten_borough/bg_rotten_borough_mood_runtime.webp", "1280x720", "118,038 bytes", "Runtime background", "Biome and title anchor", "This carried much of the perceived production value."),
    Asset("Reward Shop Counter Runtime", "assets/game/ui/menus/reward_shop_counter_runtime.webp", "1280x720", "103,602 bytes", "Runtime UI background", "Reward shop", "A single good bitmap made the shop feel more finished than its interaction depth."),
    Asset("Foxman Runtime Sheet", "assets/game/sprites/characters/foxman/char_foxman_candidate_b_prototype_sheet_alpha.png", "1774x887", "961,066 bytes", "Runtime sheet", "Atlas source", "Processed source sheet for Foxman frames."),
    Asset("Drunken Guard Runtime Sheet", "assets/game/sprites/enemies/enemy_guard_drunken_runtime_sheet_alpha.png", "2172x724", "1,162,884 bytes", "Runtime sheet", "Atlas source", "Processed source sheet for first enemy frames."),
    Asset("Tax Clerk Runtime Sheet", "assets/game/sprites/enemies/enemy_tax_clerk_runtime_sheet_alpha.png", "1774x887", "1,032,968 bytes", "Runtime sheet", "Atlas source", "Processed source sheet for bureaucratic enemies."),
    Asset("Toll Baron Runtime Sheet", "assets/game/sprites/enemies/enemy_toll_baron_runtime_sheet_alpha.png", "1536x1024", "1,818,552 bytes", "Runtime sheet", "Atlas source", "Processed source sheet for mini-boss."),
    Asset("Pickup Exit Runtime Sheet", "assets/game/props/rotten_borough/prop_pickup_exit_runtime_sheet_alpha.png", "1774x887", "1,111,434 bytes", "Runtime sheet", "Atlas source", "Processed gate and prop source sheet."),
    Asset("Rotten Borough Mood Concept", "assets/source/concepts/bg_rotten_borough_mood.png", "1774x887", "2,273,506 bytes", "Source concept", "Biome style", "The strongest world-building artifact."),
    Asset("Foxman Base Concept", "assets/source/concepts/char_foxman_base_concept_sheet.png", "1536x1024", "2,045,301 bytes", "Source concept", "Protagonist exploration", "Early style exploration."),
    Asset("Foxman Candidate B Concept", "assets/source/concepts/char_foxman_candidate_b_concept_sheet.png", "1536x1024", "2,596,673 bytes", "Source concept", "Selected protagonist direction", "Chosen as the core Foxman look."),
    Asset("Foxman Candidate C Concept", "assets/source/concepts/char_foxman_candidate_c_concept_sheet.png", "1536x1024", "2,440,707 bytes", "Source concept", "Alternate protagonist direction", "Useful comparison material."),
    Asset("Drunken Guard Concept", "assets/source/concepts/enemy_guard_drunken_concept_sheet.png", "1536x1024", "2,134,235 bytes", "Source concept", "Enemy exploration", "Strong fit for corrupt borough tone."),
    Asset("Sewer Rat Brute Concept", "assets/source/concepts/enemy_rat_sewer_brute_concept_sheet.png", "1536x1024", "2,638,685 bytes", "Source concept", "Cut enemy", "Good example of generated material that did not enter runtime scope."),
    Asset("Tax Clerk Concept", "assets/source/concepts/enemy_tax_clerk_concept_sheet.png", "1536x1024", "3,029,912 bytes", "Source concept", "Bureaucratic enemy", "Good thematic match for the corruption satire."),
    Asset("Texture Material Board", "assets/source/concepts/texture_rotten_borough_material_board.png", "1536x1024", "2,721,781 bytes", "Source concept", "Material language", "Shows grime, stone, parchment, metal, and surface direction."),
    Asset("Rotten Borough Tile Concept", "assets/source/concepts/tile_rotten_borough_concept_sheet.png", "1536x1024", "2,568,483 bytes", "Source concept", "Tile-kit direction", "Still needs conversion into runtime tile assets."),
    Asset("UI/VFX Style Board", "assets/source/concepts/ui_vfx_style_board.png", "1536x1024", "2,497,909 bytes", "Source concept", "UI and VFX language", "Useful target; runtime VFX remained partly procedural/prototype."),
    Asset("Raw Reward Shop Counter", "assets/source/ai_raw/reward_shop_counter_raw.png", "1672x941", "2,157,239 bytes", "Raw generation", "Shop background source", "Raw source for runtime shop counter."),
    Asset("Raw Reward Icon Board", "assets/source/ai_raw/reward_shop_icon_board_raw.png", "1254x1254", "2,585,482 bytes", "Raw generation", "Icon source", "Raw board behind packed reward icons."),
    Asset("Raw Foxman Runtime Chroma Sheet", "assets/source/ai_raw/char_foxman_candidate_b_prototype_sheet_chroma.png", "1774x887", "1,675,003 bytes", "Raw generation", "Runtime source", "Shows cleanup burden before alpha sheet/atlas use."),
    Asset("Raw Guard Runtime Chroma Sheet", "assets/source/ai_raw/enemy_guard_drunken_runtime_sheet_chroma.png", "2172x724", "1,764,179 bytes", "Raw generation", "Runtime source", "Shows raw-to-runtime pipeline."),
    Asset("Raw Tax Clerk Runtime Chroma Sheet", "assets/source/ai_raw/enemy_tax_clerk_runtime_sheet_chroma.png", "1774x887", "1,725,707 bytes", "Raw generation", "Runtime source", "Shows raw-to-runtime pipeline."),
    Asset("Raw Toll Baron Runtime Chroma Sheet", "assets/source/ai_raw/enemy_toll_baron_runtime_sheet_chroma.png", "1536x1024", "2,475,767 bytes", "Raw generation", "Runtime source", "Shows raw-to-runtime pipeline."),
    Asset("Raw Pickup Exit Chroma Sheet", "assets/source/ai_raw/prop_pickup_exit_runtime_sheet_chroma.png", "1774x887", "1,631,939 bytes", "Raw generation", "Runtime source", "Shows prop cleanup burden."),
]


CODE_ROWS = [
    ("Runtime config", "src/main.ts, src/style.css, src/game/GameConfig.ts, src/game/assets.ts, src/game/assetFrames.ts, src/game/smoke.ts", "Bootstraps Phaser, asset keys, display config, and smoke/manual route split."),
    ("Scenes", "BootScene, PreloadScene, TitleScene, RunScene, RewardScene, SecondRunScene, MiniBossScene, UIScene, PauseScene", "The game-shaped skeleton: title, first room, reward shop, second path, boss, HUD, pause."),
    ("Player feel", "Player.ts, PlayerMotor.ts, movement.ts, InputMapper.ts", "Movement, jumping, facing, attacks, keyboard mapping, and damage response."),
    ("Combat", "Health.ts, WeaponStats.ts, GuardEnemy.ts, HitFeedback.ts", "Health, weapon numbers, shared enemy behavior, hit text and sparks."),
    ("Progression", "RewardChoice.ts, MutationStats.ts, SkillStats.ts, ProgressStore.ts", "Rewards, mutations, active skill, local storage unlocks/kills/deaths."),
    ("Smoke tests", "tests/smoke/check-browser-routes.mjs, check-dist.mjs", "Headless browser route matrix and production artifact check."),
    ("Asset pipeline", "scripts/export-frame-atlas.py, export-runtime-background.py", "Repeatable atlas and background export pipeline."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.5, color: RGBColor = INK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Foxman Case Study | One-Shot Generation Postmortem"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "The Adventures of Foxman, a Merciless Bastard"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def preview_image_path(path: Path, image_cache: Path) -> Path:
    """Create a DOCX-friendly preview to handle WebP/alpha and constrain file size."""
    image = ImageOps.exif_transpose(Image.open(path))
    image.thumbnail((PREVIEW_MAX_EDGE, PREVIEW_MAX_EDGE), Image.Resampling.LANCZOS)
    if image.mode in ("RGBA", "LA") or (image.mode == "P" and "transparency" in image.info):
        rgba = image.convert("RGBA")
        background = Image.new("RGBA", rgba.size, (22, 19, 21, 255))
        background.alpha_composite(rgba)
        image = background.convert("RGB")
    else:
        image = image.convert("RGB")

    target = image_cache / f"{path.stem}.jpg"
    image.save(target, format="JPEG", quality=88, optimize=True)
    return target


def add_picture(doc: Document, path: Path, image_cache: Path, width) -> None:
    preview = preview_image_path(path, image_cache)
    try:
        doc.add_picture(str(preview), width=width)
    except UnrecognizedImageError:
        fallback = preview.with_suffix(".png")
        Image.open(preview).save(fallback, format="PNG", optimize=True)
        doc.add_picture(str(fallback), width=width)


def add_title(doc: Document, image_cache: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("FOXMAN")
    run.font.name = "Calibri"
    run.font.size = Pt(34)
    run.bold = True
    run.font.color.rgb = INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("A Merciless Bastard, And The Failure Mode Of One-Shot Game Generation")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("A visual case study in AI game production, agent orchestration, smoke-test mirages, and why playable is a human word.")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = MUTED

    hero = ROOT / "assets/source/concepts/char_foxman_candidate_b_concept_sheet.png"
    add_picture(doc, hero, image_cache, width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figure 1. Foxman Candidate B concept sheet, selected as the core protagonist direction.")

    add_callout(
        doc,
        "Core thesis",
        "One-shot generation produced assets, scaffolds, and plausible systems quickly. It did not produce a coherent game without orchestration, taste, manual playtesting, and quality gates.",
    )


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.italic = True
    run.font.color.rgb = MUTED


def add_callout(doc: Document, label: str, body: str, fill: str = LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_asset_block(doc: Document, asset: Asset, index: int, image_cache: Path) -> None:
    path = ROOT / asset.path
    doc.add_heading(asset.title, level=3)
    if path.exists():
        width = Inches(5.9)
        if "atlas" in asset.path or "sheet" in asset.path:
            width = Inches(6.1)
        add_picture(doc, path, image_cache, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, f"Figure {index}. {asset.title}. Path: {asset.path}")
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    rows = [
        ("Status", asset.status),
        ("Dimensions", asset.dimensions),
        ("Size", asset.size),
        ("Runtime role", asset.role),
        ("Case-study note", asset.note),
    ]
    for row_idx, (label, value) in enumerate(rows):
        left, right = table.rows[row_idx].cells
        set_cell_shading(left, HEADER)
        set_cell_text(left, label, bold=True, size=8.5, color=DARK_BLUE)
        set_cell_text(right, value, size=8.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER)
        set_cell_text(cell, header, bold=True, size=8.5, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, size=8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def build_doc() -> None:
    temp_context = TemporaryDirectory()
    image_cache = Path(temp_context.name)
    doc = Document()
    style_doc(doc)
    add_title(doc, image_cache)

    doc.add_page_break()
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The Foxman run attempted to create a Dead Cells-inspired side-scroller with generated characters, backgrounds, textures, sprite sheets, a Phaser runtime, smoke tests, gated production docs, and a postmortem-worthy project brain."
    )
    doc.add_paragraph(
        "The run produced a real scaffold: Phaser scenes, integrated assets, enemy systems, reward choices, a boss prototype, local persistence, a broad browser smoke matrix, and a large visual asset library. It did not produce a product-quality playable game in one shot."
    )
    add_callout(
        doc,
        "One-line finding",
        "The project proved that generation can create a lot of material quickly. It also proved that material is not a product.",
    )
    add_bullets(
        doc,
        [
            "The first human-facing playable link exposed a bot-driven smoke route.",
            "The first-room presentation still looked like debug scaffolding until a cleanup pass.",
            "A green completion flash washed the screen and looked like a missing-texture failure.",
            "The V1 label was too generous; the work was an internal technical slice, not a public demo.",
            "Agent orchestration, taste, and manual playtesting are production skills, not optional polish.",
        ],
    )

    doc.add_heading("2. Build Process, Prompt, Glitches, And Repair Attempts", level=1)
    doc.add_paragraph(
        "This is the core Foxman story. The prompt did not ask for a tiny prototype. It asked for a one-go, gated production run: asset design first, AI-generated textures and sprite sheets, a Phaser side-scroller, ongoing smoke tests, playtesting, and a huge initiative doc. The system answered with production-shaped output, but not production-quality judgment."
    )
    add_callout(
        doc,
        "Core process failure",
        "The run produced a convincing trail of artifacts before the first ten seconds of player experience were defensible.",
    )
    add_table(
        doc,
        ["Prompt Intent", "What The Run Produced", "Where It Broke"],
        [
            (
                "Full one-shot side-scroller with gates",
                "Initiative docs, phase reports, hopper/completed logs, ADRs, and V1 audit",
                "The docs created confidence faster than the playable experience earned it.",
            ),
            (
                "AI-generated textures, sprites, backgrounds, models",
                "Concept sheets, raw generations, alpha sheets, atlases, background art, UI boards",
                "Asset volume outpaced in-game composition, animation feel, and tile-kit polish.",
            ),
            (
                "Code bones and flesh",
                "Vite/TypeScript/Phaser scenes, player, enemies, rewards, boss, HUD, persistence",
                "Feature count grew before the first room felt authored or fun.",
            ),
            (
                "Ongoing smoke tests and playtest",
                "Broad browser smoke matrix with route completion and state assertions",
                "Bot routes were mistaken for playable links; automation hid human-control failure.",
            ),
        ],
        widths=[2.0, 2.3, 2.2],
    )

    doc.add_heading("The Glitchy Outputs", level=2)
    add_bullets(
        doc,
        [
            "Green-wash first room: a huge neon completion effect and raw-looking platforms made the room read like a missing-texture/debug screen.",
            "Autorun playable link: the shared scene was effectively a smoke route, so Foxman moved without the player controlling him.",
            "Debug scaffold aesthetic: hitboxes, projectile rectangles, procedural platform dressing, and state-driven test assumptions were still visible.",
            "Over-generous V1 label: a smoke-proven technical slice was described with language that sounded closer to demo readiness than it deserved.",
        ],
    )

    doc.add_heading("Repair Attempts", level=2)
    add_table(
        doc,
        ["Fix", "What Improved", "What Still Needed Human Production Work"],
        [
            (
                "Manual play split",
                "Smoke automation now requires smokeAuto=1; manual scene links keep player control.",
                "Controls onboarding, first-minute tutorialization, and game feel still need design passes.",
            ),
            (
                "First-room cleanup",
                "Debug state text was hidden, platforms were toned down, gate scale improved, and green-artifact smoke guard was added.",
                "The room still needs authored layout, camera composition, and real tile-kit runtime art.",
            ),
            (
                "Smoke hardening",
                "Routes now assert more state: HUD fields, hit feedback, death/restart, boss completion, reward handoff, and visual green regression.",
                "Tests still cannot judge whether the game is enjoyable, readable, or satisfying to touch.",
            ),
            (
                "Postmortem reframe",
                "The project is now positioned as an internal technical slice and production autopsy.",
                "A public demo still needs a separate human-first acceptance gate.",
            ),
        ],
        widths=[1.45, 2.45, 2.6],
    )
    add_callout(
        doc,
        "Pattern",
        "One-shot generation created output. Agent orchestration had to turn that output into evidence, reject bad milestone language, and force fixes where the player experience contradicted the reports.",
    )

    doc.add_heading("3. Token Economics", level=1)
    doc.add_paragraph(
        "The working project accounting number for the run is 5,300,000 tokens. The exact dollar value depends on the input/output split, so the correct business framing is a scenario range rather than a single pretend invoice."
    )
    add_table(
        doc,
        ["Scenario", "Input Tokens", "Output Tokens", "Estimated GPT-5.5 Standard Cost"],
        [
            ("All input", "5,300,000", "0", "$26.50"),
            ("80% input / 20% output", "4,240,000", "1,060,000", "$53.00"),
            ("70% input / 30% output", "3,710,000", "1,590,000", "$66.25"),
            ("50% input / 50% output", "2,650,000", "2,650,000", "$92.75"),
            ("All output", "0", "5,300,000", "$159.00"),
        ],
        widths=[1.9, 1.4, 1.4, 1.8],
    )
    doc.add_paragraph(
        "Pricing source: OpenAI GPT-5.5 availability and pricing note, which lists gpt-5.5 at $5 per 1M input tokens and $30 per 1M output tokens. Batch/Flex is half standard rate; Priority/Fast is 2.5x standard rate."
    )

    doc.add_heading("4. Production Timeline", level=1)
    add_numbered(
        doc,
        [
            "Initiative planning and gate design.",
            "Style lock, prompt book, and first concept generation.",
            "Vite, TypeScript, and Phaser scaffold.",
            "Movement sandbox.",
            "Combat sandbox.",
            "First room with pickup, guard, locked exit, and completion.",
            "Runtime art integration and atlas packing.",
            "Reward/shop scene and second combat path.",
            "Ranged weapon, active skill, and mutation prototypes.",
            "Toll Baron mini-boss room.",
            "HUD/readability, hit feedback, death/restart, smoke matrix.",
            "Over-generous V1 acceptance audit.",
            "Human playtest reality check and recovery fixes.",
        ],
    )

    doc.add_heading("5. Failure Modes", level=1)
    for title, body in [
        ("The Scaffold Mirage", "Scenes, entities, tests, docs, and assets made the repo look game-shaped before the game felt good."),
        ("The Smoke-Test Lie", "Automation proved route completion while hiding the first-time player experience."),
        ("The Asset Dump Problem", "Good images did not become good game art until scaled, composed, framed, and tested in motion."),
        ("The V1 Delusion", "The phrase V1 candidate sounded more finished than the build deserved."),
        ("The Taste Gap", "The agent could satisfy local instructions while missing failures a human saw in seconds."),
    ]:
        add_callout(doc, title, body)

    doc.add_heading("6. Business Lesson", level=1)
    doc.add_paragraph(
        "The useful commercial conclusion is not that AI cannot help make games. It can. The useful conclusion is that agentic production needs direction. AI lowers the cost of generating material and raises the value of selection, orchestration, QA, integration discipline, and taste."
    )
    add_bullets(
        doc,
        [
            "Use AI for fast ideation, scaffold generation, asset exploration, and test harness construction.",
            "Do not let test pass states substitute for manual playtest gates.",
            "Protect the first human experience before expanding content.",
            "Treat agent orchestration as a production discipline.",
            "Budget for review debt, not just token spend.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("7. Asset Gallery", level=1)
    doc.add_paragraph(
        "The following gallery embeds the major project assets: runtime atlases, gameplay backgrounds, runtime sheets, source concepts, and raw generations. This is the visual evidence that the run produced substantial material even when the product experience remained rough."
    )
    for idx, asset in enumerate(ASSETS, start=2):
        add_asset_block(doc, asset, idx, image_cache)

    doc.add_page_break()
    doc.add_heading("8. Code Map", level=1)
    doc.add_paragraph(
        "The codebase has real bones. The important distinction is that bones are not the same as shipped product quality."
    )
    add_table(doc, ["Area", "Files", "Role"], CODE_ROWS, widths=[1.3, 2.5, 2.7])

    doc.add_heading("9. Where The Assets Live", level=1)
    add_table(
        doc,
        ["Folder", "Contents"],
        [
            ("assets/source/concepts/", "Curated concept images used for style lock and case-study visuals."),
            ("assets/source/ai_raw/", "Raw AI generations, chroma sheets, and source boards."),
            ("assets/source/prompts/", "Prompt book for generated concepts and runtime sheets."),
            ("assets/game/sprites/", "Processed runtime sprite sheets before atlas packing."),
            ("assets/game/atlases/", "Packed Phaser atlases referenced by the game."),
            ("assets/game/backgrounds/", "Optimized runtime background art."),
            ("assets/game/ui/", "Runtime UI/menu art."),
            ("docs/case-study/", "Case-study Markdown source, asset gallery, code map, and this DOCX."),
        ],
        widths=[2.3, 4.2],
    )

    doc.add_heading("10. Recovery Plan", level=1)
    add_bullets(
        doc,
        [
            "Make the root URL the only recommended playable link.",
            "Add a polished controls overlay and first-minute tutorialization.",
            "Redesign the first room around feel, not smoke traversal.",
            "Replace procedural platform dressing with generated tile-kit runtime art.",
            "Add visual regression screenshots alongside smoke-route state checks.",
            "Rename the current artifact internally as a technical slice unless a real demo gate passes.",
        ],
    )

    doc.add_heading("Appendix: Source Documents", level=1)
    add_bullets(
        doc,
        [
            "docs/case-study/FOXMAN_CASE_STUDY.md",
            "docs/case-study/ASSET_GALLERY.md",
            "docs/case-study/CODE_MAP.md",
            "docs/FOXMANS_INITIATIVE.md",
            "docs/PHASE8_V1_ACCEPTANCE_AUDIT.md",
            "tests/smoke/check-browser-routes.mjs",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    temp_context.cleanup()
    print(OUT)


if __name__ == "__main__":
    build_doc()
